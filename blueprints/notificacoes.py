import os
from datetime import date, datetime

from flask import (
    Blueprint,
    abort,
    current_app,
    jsonify,
    redirect,
    render_template,
    request,
    send_file,
    session,
    url_for,
)

from app_core import auth as auth_core
from app_core import db as db_core
from app_core import utils as utils_core
from app_core import work_types


bp = Blueprint("notificacoes", __name__)
login_required = auth_core.login_required
STATUS_OPCOES = work_types.STATUS_OPTIONS


def _db_path():
    return current_app.config["DB_PATH"]


def _base_dir():
    return current_app.root_path


def _modelo_path():
    return os.path.join(_base_dir(), "modelo_notificacao.txt")


def _saida_dir():
    path = os.path.join(_base_dir(), "notificacoes_geradas")
    os.makedirs(path, exist_ok=True)
    return path


def get_db():
    return db_core.connect(_db_path())


def q(sql, params=()):
    return db_core.query(_db_path(), sql, params)


def q1(sql, params=()):
    return db_core.query_one(_db_path(), sql, params)


def usuario_atual():
    return auth_core.usuario_atual(q1)


def nivel_min(nivel):
    return auth_core.nivel_min(nivel, usuario_atual)


def request_int_arg(nome, default, minimo=None, maximo=None):
    return utils_core.bounded_int(request.args.get(nome), default, minimo, maximo)


def ler_modelo():
    return utils_core.ler_modelo(_modelo_path())


@bp.route("/notificacoes")
@login_required
def page():
    fs = request.args.getlist("status")
    ft = request.args.getlist("tipo")
    fl = request.args.getlist("localidade")
    fa = request.args.getlist("agente")
    d_ini = request.args.get("d_ini", "")
    d_fim = request.args.get("d_fim", "")
    busca = request.args.get("busca", "").strip()
    pagina = request_int_arg("pagina", 1, minimo=1)
    pp_str = request.args.get("por_pagina", "50")
    pp = None if pp_str == "tudo" else utils_core.safe_int(pp_str, 50)
    if pp is not None:
        pp = min(max(pp, 1), 500)
        pp_str = str(pp)

    where, params = "WHERE 1=1", []
    if d_ini:
        where += " AND f.data>=?"
        params.append(d_ini)
    if d_fim:
        where += " AND f.data<=?"
        params.append(d_fim)
    if fs:
        where += f" AND COALESCE(f.status_notificacao,'pendente') IN ({','.join('?' * len(fs))})"
        params += fs
    if ft:
        where += f" AND f.tipo_trabalho IN ({','.join('?' * len(ft))})"
        params += ft
    if fl:
        where += f" AND l.nome IN ({','.join('?' * len(fl))})"
        params += fl
    if fa:
        cond = " OR ".join(["f.agentes LIKE ?" for _ in fa])
        where += f" AND ({cond})"
        params += [f"%{a}%" for a in fa]
    if busca:
        where += (
            " AND (f.logradouro LIKE ? OR f.num_tubo LIKE ? OR f.nome_morador LIKE ? "
            "OR CAST(f.quarteirao AS TEXT) LIKE ? OR f.codigo LIKE ?)"
        )
        b = f"%{busca}%"
        params += [b, b, b, b, b]
    where += " AND f.gera_notificacao=1"

    base = (
        "SELECT f.*, l.nome AS localidade_nome FROM focos_positivos f "
        f"LEFT JOIN localidades l ON l.id_localidade=f.id_localidade {where}"
    )
    conn = get_db()
    try:
        total = conn.execute(
            "SELECT COUNT(*) FROM focos_positivos f "
            f"LEFT JOIN localidades l ON l.id_localidade=f.id_localidade {where}",
            params,
        ).fetchone()[0]

        if pp:
            total_pag = max(1, (total + pp - 1) // pp)
            pagina = min(pagina, total_pag)
            focos = conn.execute(
                base + " ORDER BY f.data DESC LIMIT ? OFFSET ?",
                params + [pp, (pagina - 1) * pp],
            ).fetchall()
        else:
            total_pag, pagina = 1, 1
            focos = conn.execute(base + " ORDER BY f.data DESC", params).fetchall()

        contadores = {}
        for row in conn.execute(
            "SELECT COALESCE(status_notificacao,'pendente') as st, COUNT(*) as cnt "
            "FROM focos_positivos WHERE gera_notificacao=1 GROUP BY st"
        ).fetchall():
            contadores[row[0]] = row[1]

        tipos_n = [
            r[0]
            for r in conn.execute(
                "SELECT DISTINCT tipo_trabalho FROM focos_positivos "
                "WHERE tipo_trabalho IS NOT NULL ORDER BY tipo_trabalho"
            ).fetchall()
        ]
        locs_n = [r[0] for r in conn.execute("SELECT DISTINCT nome FROM localidades ORDER BY nome").fetchall()]
        agentes_l = [r[0] for r in conn.execute("SELECT nome FROM agentes ORDER BY nome").fetchall()]
    finally:
        conn.close()

    return render_template(
        "notificacoes.html",
        focos=[dict(f) for f in focos],
        contadores=contadores,
        tipos=tipos_n,
        localidades_n=locs_n,
        agentes_lista=agentes_l,
        filtro_status=fs,
        filtro_tipo=ft,
        filtro_loc=fl,
        filtro_agente=fa,
        filtro_d_ini=d_ini,
        filtro_d_fim=d_fim,
        busca=busca,
        pagina=pagina,
        total_paginas=total_pag,
        total=total,
        por_pagina=pp_str,
    )


@bp.route("/notificacoes/foco/<id_foco>")
@login_required
def foco_detalhe(id_foco):
    foco = q1(
        """SELECT f.*, l.nome AS localidade_nome
           FROM focos_positivos f
           LEFT JOIN localidades l ON l.id_localidade=f.id_localidade
           WHERE f.id_foco=?""",
        (id_foco,),
    )
    if not foco:
        abort(404)
    historico = []
    if foco.get("logradouro"):
        historico = q(
            """
            SELECT v.data, v.tipo, v.visita, GROUP_CONCAT(DISTINCT a.nome) as agentes
            FROM visitas v
            LEFT JOIN visita_agentes va ON va.id_visita=v.id_visita
            LEFT JOIN agentes a ON a.id_agente=va.id_agente
            WHERE v.logradouro=? AND v.numero=?
            GROUP BY v.id_visita ORDER BY v.data DESC LIMIT 10
            """,
            (foco["logradouro"], foco.get("numero", "")),
        )
    historico_foco = q(
        """
        SELECT campo, valor_ant, valor_novo, usuario, alterado_em
        FROM focos_historico WHERE id_foco=? ORDER BY alterado_em DESC LIMIT 50
        """,
        (id_foco,),
    )
    return render_template(
        "foco_detalhe.html",
        foco=foco,
        historico=historico,
        historico_foco=historico_foco,
    )


@bp.route("/notificacoes/foco/<id_foco>/atualizar", methods=["POST"])
@login_required
@nivel_min("operador")
def foco_atualizar(id_foco):
    campos = [
        "status_notificacao",
        "tentativa_1",
        "tentativa_2",
        "tentativa_3",
        "data_entrega",
        "observacoes",
        "nome_morador",
        "logradouro",
        "numero",
        "complemento",
        "depositos",
        "agentes",
    ]
    vals = {c: request.form.get(c) or None for c in campos}
    conn = get_db()
    try:
        anterior = conn.execute(
            f"SELECT {','.join(campos)} FROM focos_positivos WHERE id_foco=?",
            (id_foco,),
        ).fetchone()

        conn.execute(
            """
            UPDATE focos_positivos SET
                status_notificacao=?,tentativa_1=?,tentativa_2=?,tentativa_3=?,
                data_entrega=?,observacoes=?,nome_morador=?,
                logradouro=?,numero=?,complemento=?,depositos=?,agentes=?
            WHERE id_foco=?
            """,
            list(vals.values()) + [id_foco],
        )

        if anterior:
            usuario = session.get("nome", "desconhecido")
            agora = datetime.now().isoformat()
            for i, campo in enumerate(campos):
                ant = anterior[i]
                nov = vals[campo]
                if str(ant or "") != str(nov or ""):
                    conn.execute(
                        """
                        INSERT INTO focos_historico
                        (id_foco, campo, valor_ant, valor_novo, usuario, alterado_em)
                        VALUES (?,?,?,?,?,?)
                        """,
                        (id_foco, campo, ant, nov, usuario, agora),
                    )

        conn.commit()
    finally:
        conn.close()
    if request.headers.get("X-Requested-With") == "XMLHttpRequest":
        return jsonify({"ok": True})
    return redirect(url_for("notificacoes.foco_detalhe", id_foco=id_foco))


@bp.route("/notificacoes/foco/<id_foco>/status", methods=["POST"])
@login_required
@nivel_min("operador")
def foco_status_rapido(id_foco):
    novo = request.json.get("status") if request.is_json else request.form.get("status")
    if novo not in STATUS_OPCOES + [None]:
        return jsonify({"erro": "Status invalido"}), 400
    conn = get_db()
    try:
        ant_row = conn.execute(
            "SELECT status_notificacao FROM focos_positivos WHERE id_foco=?",
            (id_foco,),
        ).fetchone()
        ant = ant_row[0] if ant_row else None
        conn.execute(
            "UPDATE focos_positivos SET status_notificacao=? WHERE id_foco=?",
            (novo, id_foco),
        )
        if str(ant or "") != str(novo or ""):
            conn.execute(
                """INSERT INTO focos_historico
                   (id_foco,campo,valor_ant,valor_novo,usuario,alterado_em)
                   VALUES (?,?,?,?,?,?)""",
                (
                    id_foco,
                    "status_notificacao",
                    ant,
                    novo,
                    session.get("nome", "desconhecido"),
                    datetime.now().isoformat(),
                ),
            )
        conn.commit()
    finally:
        conn.close()
    return jsonify({"ok": True, "status": novo})


@bp.route("/notificacoes/imprimir", methods=["POST"])
@login_required
@nivel_min("operador")
def imprimir():
    ids = request.form.getlist("ids")
    if not ids:
        return redirect(url_for("notificacoes.page"))
    conn = get_db()
    focos = []
    for id_foco in ids:
        row = conn.execute(
            """
            SELECT f.*, l.nome AS localidade_nome FROM focos_positivos f
            LEFT JOIN localidades l ON l.id_localidade=f.id_localidade
            WHERE f.id_foco=? AND f.gera_notificacao=1
            """,
            (id_foco,),
        ).fetchone()
        if row:
            focos.append(dict(row))
    if not focos:
        conn.close()
        return "Nenhum foco valido.", 400
    try:
        caminho = gerar_docx(focos)
    except Exception as exc:
        conn.close()
        return f"Erro ao gerar DOCX: {exc}", 500
    for foco in focos:
        conn.execute(
            """UPDATE focos_positivos SET status_notificacao='impressa'
               WHERE id_foco=? AND COALESCE(status_notificacao,'pendente')='pendente'""",
            (foco["id_foco"],),
        )
    conn.commit()
    conn.close()
    return send_file(
        caminho,
        as_attachment=True,
        download_name=f"notificacoes_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
    )


def formatar_data_br(data_iso):
    meses = [
        "janeiro",
        "fevereiro",
        "marco",
        "abril",
        "maio",
        "junho",
        "julho",
        "agosto",
        "setembro",
        "outubro",
        "novembro",
        "dezembro",
    ]
    try:
        d = datetime.strptime(str(data_iso)[:10], "%Y-%m-%d")
        return f"{d.day} de {meses[d.month - 1]} de {d.year}"
    except (ValueError, TypeError, AttributeError):
        return "______"


def remover_bordas(tabela):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    for row in tabela.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for side in ["top", "bottom", "left", "right", "insideH", "insideV"]:
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "none")
                borders.append(border)
            tcPr.append(borders)


def gerar_via(doc, foco, modelo):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    def par(
        texto="",
        bold=False,
        italic=False,
        size=11,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        sb=0,
        sa=6,
    ):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.space_after = Pt(sa)
        if texto:
            r = p.add_run(texto)
            r.bold = bold
            r.italic = italic
            r.font.size = Pt(size)
        return p

    def par_mixed(partes, align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.space_after = Pt(sa)
        for texto, bold, italic, size in partes:
            r = p.add_run(texto)
            r.bold = bold
            r.italic = italic
            r.font.size = Pt(size)
        return p

    cab = doc.add_table(rows=1, cols=3)
    cab.style = "Table Grid"
    remover_bordas(cab)

    cl = cab.rows[0].cells[0]
    cl.width = Cm(3.5)
    pl = cl.paragraphs[0]
    pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_pref = os.path.join(_base_dir(), "static", "img", "logo_prefeitura.png")
    try:
        pl.add_run().add_picture(logo_pref, width=Cm(3))
    except Exception:
        pl.add_run("[LOGO PREFEITURA]")

    ct = cab.rows[0].cells[1]
    for i, linha in enumerate(modelo.get("CABECALHO", "").split("\n")):
        p = ct.paragraphs[0] if i == 0 else ct.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(linha)
        r.bold = i == 0
        r.font.size = Pt(10 if i == 0 else 9)

    cr = cab.rows[0].cells[2]
    cr.width = Cm(3.5)
    pr = cr.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_end = os.path.join(_base_dir(), "static", "img", "logo_endemias.png")
    try:
        pr.add_run().add_picture(logo_end, width=Cm(3))
    except Exception:
        pr.add_run("[LOGO ENDEMIAS]")

    tblPr = cab._tbl.tblPr
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    tblPr.append(jc)

    doc.add_paragraph()
    par(
        f"Almirante Tamandare (PR), ______/______ de {date.today().year}.",
        align=WD_ALIGN_PARAGRAPH.RIGHT,
        size=10,
        sa=8,
    )
    par(
        modelo.get("TITULO", "COMUNICADO / NOTIFICACAO"),
        bold=True,
        size=13,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        sa=4,
    )
    par(
        modelo.get("SAUDACAO", "Prezado(a) Senhor(a) PROPRIETARIO/RESPONSAVEL"),
        bold=True,
        size=11,
        sa=8,
    )

    end_fmt = f"{foco.get('logradouro') or ''}, {foco.get('numero') or 's/n'}".strip(", ")
    loc_fmt = foco.get("localidade_nome") or foco.get("localidade") or ""
    qrt_fmt = f"Quarteirao {foco.get('quarteirao')}" if foco.get("quarteirao") else ""
    loc_linha = " - ".join(filter(None, [loc_fmt, qrt_fmt]))

    corpo = (
        modelo.get("CORPO", "")
        .replace("{endereco}", end_fmt)
        .replace("{localidade}", loc_linha)
        .replace("{data_visita}", formatar_data_br(foco.get("data")))
    )
    partes = []
    marcador = "Aedes aegypti"
    while marcador in corpo:
        idx = corpo.index(marcador)
        if corpo[:idx]:
            partes.append((corpo[:idx], False, False, 11))
        partes.append((marcador, False, True, 11))
        corpo = corpo[idx + len(marcador) :]
    if corpo:
        partes.append((corpo, False, False, 11))
    par_mixed(partes, sa=8)

    par(modelo.get("AVISO", ""), bold=True, size=11, sa=8)
    par(modelo.get("CONTATO", ""), size=10, sa=10)
    campos = [
        ("LOCALIDADE", loc_linha or "___________________________"),
        ("ENDERECO", end_fmt or "___________________________"),
        ("MORADOR", foco.get("nome_morador") or "___________________________"),
        ("DEPOSITO(S)", foco.get("depositos") or "___________________________"),
        ("AGENTE(S)", foco.get("agentes") or "___________________________"),
    ]
    if foco.get("observacoes"):
        campos.append(("OBSERVACOES", foco["observacoes"]))
    for label, valor in campos:
        par_mixed([(f"- {label}: ", True, True, 11), (valor, False, True, 11)], sa=3)

    doc.add_paragraph()
    ass = doc.add_table(rows=2, cols=2)
    ass.style = "Table Grid"
    remover_bordas(ass)
    for i, txt in enumerate(["_" * 35, "_" * 35]):
        cell = ass.rows[0].cells[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].add_run(txt)
    for i, txt in enumerate(["Vigilancia Ambiental / Setor de Endemias", "Proprietario / Responsavel"]):
        cell = ass.rows[1].cells[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].add_run(txt).bold = True

    doc.add_paragraph()
    par(modelo.get("RODAPE", ""), size=8, align=WD_ALIGN_PARAGRAPH.CENTER, sa=0)
    if foco.get("codigo"):
        par(
            f"No da notificacao: {foco['codigo']}",
            size=7,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            sa=0,
        )


def gerar_docx(focos):
    from docx import Document
    from docx.shared import Cm

    modelo = ler_modelo()
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(1.5)
        section.left_margin = section.right_margin = Cm(2)
    for i, foco in enumerate(focos):
        gerar_via(doc, foco, modelo)
        if i < len(focos) - 1:
            doc.add_page_break()
    caminho = os.path.join(_saida_dir(), f"notificacoes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    doc.save(caminho)
    return caminho


def _focos_para_impressao(ids):
    conn = get_db()
    focos = []
    try:
        for id_foco in ids:
            row = conn.execute(
                """
                SELECT f.*, l.nome AS localidade_nome
                FROM focos_positivos f
                LEFT JOIN localidades l ON l.id_localidade = f.id_localidade
                WHERE f.id_foco = ? AND f.gera_notificacao = 1
                """,
                (id_foco,),
            ).fetchone()
            if row:
                focos.append(dict(row))
    finally:
        conn.close()
    return focos


@bp.route("/notificacoes/foco/<id_foco>/imprimir-html", methods=["POST"])
@login_required
@nivel_min("operador")
def imprimir_html_single(id_foco):
    focos = _focos_para_impressao([id_foco])
    if not focos:
        abort(404)
    conn = get_db()
    try:
        conn.execute(
            """UPDATE focos_positivos SET status_notificacao='impressa'
               WHERE id_foco=? AND COALESCE(status_notificacao,'pendente')='pendente'""",
            (id_foco,),
        )
        conn.commit()
    finally:
        conn.close()
    return render_template(
        "notificacao_print.html",
        focos=focos,
        auto_print=True,
        modelo=type("M", (), ler_modelo())(),
    )


@bp.route("/notificacoes/imprimir-html", methods=["POST"])
@login_required
@nivel_min("operador")
def imprimir_html_lote():
    ids = request.form.getlist("ids")
    if not ids:
        return redirect(url_for("notificacoes.page"))
    focos = _focos_para_impressao(ids)
    if not focos:
        return "Nenhum foco valido.", 400
    conn = get_db()
    try:
        for foco in focos:
            conn.execute(
                """UPDATE focos_positivos SET status_notificacao='impressa'
                   WHERE id_foco=? AND COALESCE(status_notificacao,'pendente')='pendente'""",
                (foco["id_foco"],),
            )
        conn.commit()
    finally:
        conn.close()
    return render_template(
        "notificacao_print.html",
        focos=focos,
        auto_print=True,
        modelo=type("M", (), ler_modelo())(),
    )
