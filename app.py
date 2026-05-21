"""
Endemias â€” Sistema de GestÃ£o Integrado  v3
Setor de Endemias / VigilÃ¢ncia Ambiental â€” Almirante TamandarÃ©-PR

Servidor Ãºnico: rode em um computador e os demais acessam via http://IP:5000
"""
import os, sqlite3, json
import logging
import logging.handlers
from datetime import date, datetime, timedelta

from flask import (Flask, render_template, request, redirect, url_for,
                   send_file, jsonify, abort, session)
from flask_wtf.csrf import CSRFProtect, CSRFError

from app_core import auth as auth_core
from app_core import db as db_core
from app_core import import_history
from app_core import modules as modules_core
from app_core import uploads as uploads_core
from app_core import utils as utils_core
from app_core import version as version_core
from app_core import work_types
from blueprints.admin import bp as admin_bp
from blueprints.agenda import bp as agenda_bp
from blueprints.conta_ovos_sispncd import bp as conta_ovos_sispncd_bp
from blueprints.consultas import bp as consultas_bp
from blueprints.esporotricose import bp as esporotricose_bp
from blueprints.exportacoes import bp as exportacoes_bp
from blueprints.mapa import bp as mapa_bp
from blueprints.processar import bp as processar_bp
from blueprints.relatorio_agente import bp as relatorio_agente_bp

# â”€â”€ ValidaÃ§Ã£o de upload de arquivos (SEC-04) â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
# xlsx Ã© um ZIP internamente â€” assinatura PK\x03\x04 nos primeiros bytes
def _validar_arquivo_xlsx(file_storage):
    return uploads_core.validar_arquivo_xlsx(file_storage)

BASE_DIR    = os.path.dirname(os.path.abspath(__file__))
DB_PATH     = os.path.join(BASE_DIR, "endemias.db")
CONFIG_PATH = os.path.join(BASE_DIR, "config.json")
MODELO_PATH = os.path.join(BASE_DIR, "modelo_notificacao.txt")
SAIDA_DIR   = os.path.join(BASE_DIR, "notificacoes_geradas")
UPLOAD_TEMP = os.path.join(BASE_DIR, "uploads_temp")

os.makedirs(SAIDA_DIR, exist_ok=True)
os.makedirs(UPLOAD_TEMP, exist_ok=True)

app = Flask(__name__)
app.config["DB_PATH"] = DB_PATH
app.config["CONFIG_PATH"] = CONFIG_PATH
app.config["UPLOAD_TEMP"] = UPLOAD_TEMP

# â”€â”€ Logging estruturado em arquivo rotativo â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
_log_path = os.path.join(BASE_DIR, "endemias.log")
_log_handler = logging.handlers.RotatingFileHandler(
    _log_path, maxBytes=5 * 1024 * 1024, backupCount=3, encoding="utf-8"
)
_log_handler.setFormatter(logging.Formatter(
    "%(asctime)s [%(levelname)s] %(funcName)s: %(message)s"
))
logging.getLogger().addHandler(_log_handler)
logging.getLogger().setLevel(logging.WARNING)

# â”€â”€ Secret key: lida de arquivo local, gerada automaticamente se nÃ£o existir â”€â”€
_KEY_FILE = os.path.join(BASE_DIR, "secret.key")
if os.path.exists(_KEY_FILE):
    with open(_KEY_FILE, "rb") as _f:
        app.secret_key = _f.read()
else:
    import secrets as _secrets
    _k = _secrets.token_bytes(32)
    with open(_KEY_FILE, "wb") as _f:
        _f.write(_k)
    app.secret_key = _k
    print("[OK] secret.key gerado. Nunca compartilhe ou versione este arquivo.")

# â”€â”€ ConfiguraÃ§Ãµes de sessÃ£o segura â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
app.config["PERMANENT_SESSION_LIFETIME"] = timedelta(hours=8)
app.config["SESSION_COOKIE_HTTPONLY"]    = True
app.config["SESSION_COOKIE_SAMESITE"]   = "Lax"

# Limite simples de tentativas de login por IP+usuario.
# Como o sistema roda em processo unico na rede local, memoria atende bem sem nova dependencia.
LOGIN_MAX_TENTATIVAS = auth_core.LOGIN_MAX_TENTATIVAS
LOGIN_JANELA_SEG     = auth_core.LOGIN_JANELA_SEG
_login_tentativas    = auth_core.login_tentativas

# â”€â”€ CSRF Protection (SEC-03) â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
# Protege todos os formulÃ¡rios POST contra ataques de cross-site request forgery.
# Rotas de SSE (stream) sÃ£o isentas automaticamente (mÃ©todo GET).
# Rotas de API JSON que recebem o header X-CSRFToken tambÃ©m sÃ£o validadas.
app.config["WTF_CSRF_TIME_LIMIT"]   = 3600  # token vÃ¡lido por 1h
app.config["WTF_CSRF_CHECK_DEFAULT"] = True
csrf = CSRFProtect(app)
app.register_blueprint(admin_bp)

# Isentar rotas que nÃ£o precisam de CSRF (SSE â€” usam GET, sem estado)
# Nota: rotas GET nÃ£o sÃ£o afetadas pelo CSRF de qualquer forma.
# As Ãºnicas isenÃ§Ãµes necessÃ¡rias sÃ£o endpoints chamados por sistemas externos.
# Por ora, nenhuma isenÃ§Ã£o â€” todos os POST sÃ£o protegidos.

@app.errorhandler(CSRFError)
def handle_csrf_error(e):
    """Retorna erro amigÃ¡vel quando o token CSRF falha ou expira."""
    logging.warning(f"CSRFError: {e.description} | IP: {request.remote_addr} | URL: {request.url}")
    if request.is_json or request.headers.get("X-Requested-With") == "XMLHttpRequest":
        return jsonify({"erro": "Token de seguranÃ§a expirado. Recarregue a pÃ¡gina e tente novamente."}), 400
    return render_template("erro_csrf.html"), 400


@app.template_filter("data_br")
def filtro_data_br(valor):
    meses = ["janeiro","fevereiro","marÃ§o","abril","maio","junho",
             "julho","agosto","setembro","outubro","novembro","dezembro"]
    try:
        from datetime import datetime as _dt
        d = _dt.strptime(str(valor)[:10], "%Y-%m-%d")
        return f"{d.day} de {meses[d.month-1]} de {d.year}"
    except (ValueError, TypeError, AttributeError):
        return str(valor) if valor else "______"

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  CONSTANTES UI
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

STATUS_OPCOES = work_types.STATUS_OPTIONS
STATUS_CORES = work_types.STATUS_COLORS
TIPO_CORES = work_types.WORK_TYPE_COLORS
TIPO_LABELS = work_types.WORK_TYPE_LABELS
TIPOS_TRABALHO = work_types.WORK_TYPES
AGENDA_TIPO_COR = work_types.AGENDA_TYPE_COLORS
AGENDA_TIPO_LABEL = work_types.AGENDA_TYPE_LABELS
AGENDA_TIPOS = work_types.AGENDA_TYPES
AGENDA_FORM_LABEL = work_types.AGENDA_FORM_LABELS

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  BANCO
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def get_db():
    return db_core.connect(DB_PATH)

def q(sql, params=()):
    return db_core.query(DB_PATH, sql, params)

def q1(sql, params=()):
    return db_core.query_one(DB_PATH, sql, params)

def qval(sql, params=()):
    return db_core.scalar(DB_PATH, sql, params)

# FIX ARQ-01: Cache simples com TTL para evitar queries repetidas em todo request
# Localidades e agentes mudam raramente â€” cache de 60s Ã© seguro
import time as _time
_glob_cache: dict = {}
_CACHE_TTL = 60  # segundos

def _cached_q(key, sql, params=()):
    now = _time.monotonic()
    if key not in _glob_cache or now - _glob_cache[key][0] > _CACHE_TTL:
        _glob_cache[key] = (now, q(sql, params))
    return _glob_cache[key][1]

def invalidar_cache_globals():
    """Chamar apÃ³s ETL ou apÃ³s criar/editar agentes e localidades."""
    _glob_cache.clear()

app.extensions["invalidar_cache_globals"] = invalidar_cache_globals
app.register_blueprint(agenda_bp)
app.register_blueprint(conta_ovos_sispncd_bp)
app.register_blueprint(consultas_bp)
app.register_blueprint(esporotricose_bp)
app.register_blueprint(exportacoes_bp)
app.register_blueprint(mapa_bp)
app.register_blueprint(processar_bp)
app.register_blueprint(relatorio_agente_bp)

def garantir_tabela_importacoes(conn=None):
    return import_history.garantir_tabela_importacoes(get_db, conn)

def registrar_importacao(job_id, arquivos, status="upload", usuario=None):
    usuario = usuario or session.get("nome", "")
    return import_history.registrar_importacao(get_db, job_id, arquivos, status, usuario)

def atualizar_importacao(job_id, status, dry_run_ok=None, commit_ok=None, sumario=None, erro=None):
    return import_history.atualizar_importacao(
        get_db, job_id, status, dry_run_ok=dry_run_ok,
        commit_ok=commit_ok, sumario=sumario, erro=erro,
    )

def listar_importacoes_recentes(limite=10):
    return import_history.listar_importacoes_recentes(get_db, limite)

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  AUTH
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

# â”€â”€ FunÃ§Ãµes de hash de senha â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
# COMPATIBILIDADE: hashes antigos (SHA-256 puro) ainda funcionam para login,
# mas ao salvar nova senha sempre usa pbkdf2:sha256 com salt.

def _hash_legado(senha):
    """SHA-256 sem salt â€” usado APENAS para verificar hashes antigos."""
    return auth_core.hash_legado(senha)

def _hash(senha):
    """Gera hash seguro com pbkdf2:sha256 e salt aleatÃ³rio (werkzeug)."""
    return auth_core.hash_senha(senha)

def _verificar_senha(senha_digitada, hash_armazenado):
    """
    Verifica senha contra hash armazenado.
    Aceita tanto hashes werkzeug (pbkdf2:sha256:...) quanto hashes legados (SHA-256 puro).
    Ao autenticar com hash legado, atualiza automaticamente para hash seguro.
    Retorna (ok: bool, novo_hash: str|None) â€” novo_hash != None significa que deve ser salvo.
    """
    return auth_core.verificar_senha(senha_digitada, hash_armazenado)
        # Hash moderno werkzeug
        # Hash legado SHA-256 â€” verificar e fazer upgrade

def usuario_atual():
    return auth_core.usuario_atual(q1)

def login_required(f):
    return auth_core.login_required(f)

def nivel_min(nivel):
    """Decorator: exige nÃ­vel mÃ­nimo (admin > operador > visualizador)."""
    return auth_core.nivel_min(nivel, usuario_atual)

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  HELPERS
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def hoje():             return utils_core.hoje()
def data_n_dias(n=30):  return utils_core.data_n_dias(n)
def data_ano():         return utils_core.data_ano()

def safe_int(v, default=0):
    return utils_core.safe_int(v, default)

def request_int_arg(nome, default, minimo=None, maximo=None):
    return utils_core.bounded_int(request.args.get(nome), default, minimo, maximo)

def build_where(params_dict, alias_v="v", alias_l="l", alias_a="a"):
    where, params = "WHERE 1=1", []
    d_ini = params_dict.get("d_ini") or data_n_dias(365)
    d_fim = params_dict.get("d_fim") or hoje()
    where += f" AND {alias_v}.data BETWEEN ? AND ?"; params += [d_ini, d_fim]

    def getlist(k):
        if hasattr(params_dict, "getlist"): return params_dict.getlist(k)
        v = params_dict.get(k, [])
        return v if isinstance(v, list) else [v]

    tipos = getlist("tipo"); locs = getlist("localidade"); ags = getlist("agente")

    if tipos:
        where += f" AND {alias_v}.tipo IN ({','.join('?'*len(tipos))})"; params += tipos
    if locs:
        where += f" AND {alias_l}.nome IN ({','.join('?'*len(locs))})"; params += locs
    if ags:
        cond = " OR ".join([
            f"EXISTS(SELECT 1 FROM visita_agentes va2 JOIN agentes a2 ON a2.id_agente=va2.id_agente "
            f"WHERE va2.id_visita={alias_v}.id_visita AND a2.nome=?)"
            for _ in ags
        ])
        where += f" AND ({cond})"; params += ags

    return where, params

def ler_modelo():
    return utils_core.ler_modelo(MODELO_PATH)

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  CONTEXTO GLOBAL
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

@app.context_processor
def inject_globals():
    # FIX ARQ-01: localidades, agentes e tipos_v sÃ£o cacheados por 60s
    # pendentes sempre Ã© consultado em real-time (muda com frequÃªncia)
    localidades = _cached_q("localidades", "SELECT nome FROM localidades ORDER BY nome")
    agentes     = _cached_q("agentes",     "SELECT nome FROM agentes ORDER BY nome")
    tipos_v     = _cached_q("tipos_visita","SELECT DISTINCT tipo FROM visitas WHERE tipo IS NOT NULL ORDER BY tipo")
    pendentes   = qval("SELECT COUNT(*) FROM focos_positivos WHERE status_notificacao='pendente' AND gera_notificacao=1")
    u           = usuario_atual()
    return dict(
        localidades_glob=[r["nome"] for r in localidades],
        agentes_glob    =[r["nome"] for r in agentes],
        tipos_glob      =[r["tipo"] for r in tipos_v],
        nav_pendentes   =pendentes,
        hoje            =hoje(),
        STATUS_OPCOES   =STATUS_OPCOES,
        STATUS_CORES    =STATUS_CORES,
        TIPO_CORES      =TIPO_CORES,
        TIPO_LABELS     =TIPO_LABELS,
        TIPOS_TRABALHO  =TIPOS_TRABALHO,
        AGENDA_TIPO_CORES=AGENDA_TIPO_COR,
        AGENDA_TIPO_LABELS=AGENDA_TIPO_LABEL,
        AGENDA_TIPOS    =AGENDA_TIPOS,
        AGENDA_FORM_LABELS=AGENDA_FORM_LABEL,
        usuario_atual   =u,
        topbar_modules  = modules_core.visible_modules(u, area="topbar"),
        sidebar_groups  = modules_core.grouped_modules(modules_core.visible_modules(u, area="sidebar")),
        home_modules    = modules_core.visible_modules(u, area="home"),
        APP_VERSION      = version_core.APP_VERSION,
        APP_VERSION_DATE = version_core.APP_VERSION_DATE,
        APP_VERSION_LABEL= version_core.APP_VERSION_LABEL,
    )

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  ROTAS â€” AUTH
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def _url_segura(target):
    """Retorna True se a URL alvo Ã© do prÃ³prio servidor (previne open redirect)."""
    return auth_core.url_segura(target)

def _chave_login(usuario):
    return auth_core.chave_login(usuario)

def _login_bloqueado(chave, agora=None):
    return auth_core.login_bloqueado(chave, agora)

def _registrar_login_falha(chave, agora=None):
    return auth_core.registrar_login_falha(chave, agora)

def _limpar_login_falhas(chave):
    return auth_core.limpar_login_falhas(chave)


@app.route("/login", methods=["GET", "POST"])
@csrf.exempt   # login nÃ£o tem sessÃ£o prÃ©via para validar token â€” protegido pelo rate limit da senha
def login():
    if session.get("uid"):
        return redirect(url_for("home"))
    erro = None
    if request.method == "POST":
        usuario = request.form.get("usuario", "").strip()
        senha   = request.form.get("senha", "")
        chave   = _chave_login(usuario)
        if _login_bloqueado(chave):
            logging.warning(f"Login bloqueado por excesso de tentativas | usuario={usuario} | IP={request.remote_addr}")
            return render_template("login.html", erro="Muitas tentativas incorretas. Aguarde alguns minutos e tente novamente."), 429
        u = q1("SELECT * FROM usuarios WHERE usuario=? AND ativo=1", (usuario,))
        if u:
            ok, novo_hash = _verificar_senha(senha, u["senha_hash"])
            if ok:
                _limpar_login_falhas(chave)
                session.permanent = True          # respeita PERMANENT_SESSION_LIFETIME
                session["uid"]    = u["id_usuario"]
                session["nivel"]  = u["nivel"]
                session["nome"]   = u["nome"]
                # Upgrade silencioso de hash legado para pbkdf2
                if novo_hash:
                    try:
                        conn = get_db()
                        conn.execute("UPDATE usuarios SET senha_hash=? WHERE id_usuario=?",
                                     (novo_hash, u["id_usuario"]))
                        conn.commit()
                        conn.close()
                    except Exception:
                        pass
                # Redirecionar com seguranÃ§a (sem open redirect)
                dest = request.args.get("next", "")
                if not dest or not _url_segura(dest):
                    dest = url_for("home")
                return redirect(dest)
        erro = "UsuÃ¡rio ou senha incorretos."
        _registrar_login_falha(chave)
    return render_template("login.html", erro=erro)

@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))

@app.route("/minha-senha", methods=["GET", "POST"])
@login_required
def minha_senha():
    erro = ok = None
    if request.method == "POST":
        atual = request.form.get("atual", "")
        nova  = request.form.get("nova", "")
        conf  = request.form.get("confirmar", "")
        u = q1("SELECT * FROM usuarios WHERE id_usuario=?", (session["uid"],))
        senha_ok, _ = _verificar_senha(atual, u["senha_hash"])
        if not senha_ok:
            erro = "Senha atual incorreta."
        elif len(nova) < 6:
            erro = "A nova senha deve ter ao menos 6 caracteres."
        elif nova != conf:
            erro = "As senhas nÃ£o coincidem."
        else:
            conn = get_db()
            conn.execute("UPDATE usuarios SET senha_hash=? WHERE id_usuario=?",
                         (_hash(nova), session["uid"]))
            conn.commit(); conn.close()
            ok = "Senha alterada com sucesso."
    return render_template("minha_senha.html", erro=erro, ok=ok)

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  ROTAS â€” PÃGINAS PRINCIPAIS
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

@app.route("/")
@login_required
def home():
    conn = get_db()
    try:
        ano_ini  = data_ano()
        kpis = {
            "visitas_hoje":     conn.execute("SELECT COUNT(*) FROM visitas WHERE data=?", (hoje(),)).fetchone()[0],
            "visitas_ano":      conn.execute("SELECT COUNT(*) FROM visitas WHERE data>=?", (ano_ini,)).fetchone()[0],
            "focos_pendentes":  conn.execute("SELECT COUNT(*) FROM focos_positivos WHERE status_notificacao='pendente' AND gera_notificacao=1").fetchone()[0],
            "agentes_ativos":   conn.execute("SELECT COUNT(DISTINCT id_agente) FROM visita_agentes va JOIN visitas v ON v.id_visita=va.id_visita WHERE v.data>=?", (data_n_dias(30),)).fetchone()[0],
            "coletas_total":    conn.execute("SELECT COUNT(*) FROM coletas").fetchone()[0],
            "positivos_aeg":    conn.execute("SELECT COUNT(*) FROM resultados_laboratorio WHERE aegypt_larvas>0 OR aegypt_pupas>0 OR aegypt_exuvias>0 OR aegypt_adulto>0").fetchone()[0],
        }
        atividade   = conn.execute("SELECT data, COUNT(*) as total FROM visitas WHERE data>=? GROUP BY data ORDER BY data DESC LIMIT 14", (data_n_dias(14),)).fetchall()
        dist_tipo   = conn.execute("SELECT tipo, COUNT(*) as total FROM visitas WHERE data>=? GROUP BY tipo ORDER BY total DESC", (ano_ini,)).fetchall()
        focos_rec   = conn.execute("""
            SELECT f.*, l.nome as localidade_nome FROM focos_positivos f
            LEFT JOIN localidades l ON l.id_localidade=f.id_localidade
            WHERE f.gera_notificacao=1 ORDER BY f.processado_em DESC LIMIT 5
        """).fetchall()
    finally:
        conn.close()
    return render_template("home.html",
        kpis=kpis,
        atividade=[dict(r) for r in atividade],
        dist_tipo=[dict(r) for r in dist_tipo],
        focos_recentes=[dict(r) for r in focos_rec],
    )

# â”€â”€ COD-03: lÃ³gica de dados do relatÃ³rio de agente extraÃ­da aqui â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
# Usada tanto pela rota PDF quanto pela API, sem duplicaÃ§Ã£o.
# â”€â”€ NOTIFICAÃ‡Ã•ES â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

@app.route("/notificacoes")
@login_required
def notificacoes():
    fs = request.args.getlist("status")
    ft = request.args.getlist("tipo")
    fl = request.args.getlist("localidade")
    fa = request.args.getlist("agente")
    d_ini   = request.args.get("d_ini", "")
    d_fim   = request.args.get("d_fim", "")
    busca   = request.args.get("busca", "").strip()
    pagina  = request_int_arg("pagina", 1, minimo=1)
    pp_str  = request.args.get("por_pagina", "50")
    pp      = None if pp_str == "tudo" else safe_int(pp_str, 50)
    if pp is not None:
        pp = min(max(pp, 1), 500)
        pp_str = str(pp)

    where, params = "WHERE 1=1", []
    if d_ini:   where += " AND f.data>=?"; params.append(d_ini)
    if d_fim:   where += " AND f.data<=?"; params.append(d_fim)
    if fs:      where += f" AND COALESCE(f.status_notificacao,'pendente') IN ({','.join('?'*len(fs))})"; params+=fs
    if ft:      where += f" AND f.tipo_trabalho IN ({','.join('?'*len(ft))})"; params+=ft
    if fl:      where += f" AND l.nome IN ({','.join('?'*len(fl))})"; params+=fl
    if fa:
        cond = " OR ".join(["f.agentes LIKE ?" for _ in fa])
        where += f" AND ({cond})"; params += [f"%{a}%" for a in fa]
    if busca:
        where += " AND (f.logradouro LIKE ? OR f.num_tubo LIKE ? OR f.nome_morador LIKE ? OR CAST(f.quarteirao AS TEXT) LIKE ? OR f.codigo LIKE ?)"
        b = f"%{busca}%"; params += [b, b, b, b, b]
    where += " AND f.gera_notificacao=1"

    base = f"SELECT f.*, l.nome AS localidade_nome FROM focos_positivos f LEFT JOIN localidades l ON l.id_localidade=f.id_localidade {where}"
    conn = get_db()
    total = conn.execute(f"SELECT COUNT(*) FROM focos_positivos f LEFT JOIN localidades l ON l.id_localidade=f.id_localidade {where}", params).fetchone()[0]

    if pp:
        total_pag = max(1, (total + pp - 1) // pp)
        pagina    = min(pagina, total_pag)
        focos     = conn.execute(base + " ORDER BY f.data DESC LIMIT ? OFFSET ?", params + [pp, (pagina-1)*pp]).fetchall()
    else:
        total_pag, pagina = 1, 1
        focos = conn.execute(base + " ORDER BY f.data DESC", params).fetchall()

    contadores = {}
    for row in conn.execute("SELECT COALESCE(status_notificacao,'pendente') as st, COUNT(*) as cnt FROM focos_positivos WHERE gera_notificacao=1 GROUP BY st").fetchall():
        contadores[row[0]] = row[1]

    tipos_n   = [r[0] for r in conn.execute("SELECT DISTINCT tipo_trabalho FROM focos_positivos WHERE tipo_trabalho IS NOT NULL ORDER BY tipo_trabalho").fetchall()]
    locs_n    = [r[0] for r in conn.execute("SELECT DISTINCT nome FROM localidades ORDER BY nome").fetchall()]
    agentes_l = [r[0] for r in conn.execute("SELECT nome FROM agentes ORDER BY nome").fetchall()]
    conn.close()

    return render_template("notificacoes.html",
        focos=[dict(f) for f in focos], contadores=contadores,
        tipos=tipos_n, localidades_n=locs_n, agentes_lista=agentes_l,
        filtro_status=fs, filtro_tipo=ft, filtro_loc=fl, filtro_agente=fa,
        filtro_d_ini=d_ini, filtro_d_fim=d_fim, busca=busca,
        pagina=pagina, total_paginas=total_pag, total=total, por_pagina=pp_str,
    )

@app.route("/notificacoes/foco/<id_foco>")
@login_required
def foco_detalhe(id_foco):
    foco = q1("SELECT f.*, l.nome AS localidade_nome FROM focos_positivos f LEFT JOIN localidades l ON l.id_localidade=f.id_localidade WHERE f.id_foco=?", (id_foco,))
    if not foco: abort(404)
    historico = []
    if foco.get("logradouro"):
        historico = q("""
            SELECT v.data, v.tipo, v.visita, GROUP_CONCAT(DISTINCT a.nome) as agentes
            FROM visitas v
            LEFT JOIN visita_agentes va ON va.id_visita=v.id_visita
            LEFT JOIN agentes a ON a.id_agente=va.id_agente
            WHERE v.logradouro=? AND v.numero=?
            GROUP BY v.id_visita ORDER BY v.data DESC LIMIT 10
        """, (foco["logradouro"], foco.get("numero", "")))
    historico_foco = q("""
        SELECT campo, valor_ant, valor_novo, usuario, alterado_em
        FROM focos_historico WHERE id_foco=? ORDER BY alterado_em DESC LIMIT 50
    """, (id_foco,))
    return render_template("foco_detalhe.html", foco=foco, historico=historico,
                           historico_foco=historico_foco)

@app.route("/notificacoes/foco/<id_foco>/atualizar", methods=["POST"])
@login_required
@nivel_min("operador")
def foco_atualizar(id_foco):
    campos = ["status_notificacao","tentativa_1","tentativa_2","tentativa_3",
              "data_entrega","observacoes","nome_morador","logradouro","numero","complemento","depositos","agentes"]
    vals   = {c: request.form.get(c) or None for c in campos}
    conn   = get_db()
    try:
        # Ler valores anteriores para auditoria
        anterior = conn.execute(
            f"SELECT {','.join(campos)} FROM focos_positivos WHERE id_foco=?", (id_foco,)
        ).fetchone()

        conn.execute("""
            UPDATE focos_positivos SET
                status_notificacao=?,tentativa_1=?,tentativa_2=?,tentativa_3=?,
                data_entrega=?,observacoes=?,nome_morador=?,
                logradouro=?,numero=?,complemento=?,depositos=?,agentes=?
            WHERE id_foco=?
        """, list(vals.values()) + [id_foco])

        # Registrar apenas campos que mudaram
        if anterior:
            usuario = session.get("nome", "desconhecido")
            agora   = datetime.now().isoformat()
            for i, campo in enumerate(campos):
                ant = anterior[i]
                nov = vals[campo]
                if str(ant or "") != str(nov or ""):
                    conn.execute("""
                        INSERT INTO focos_historico (id_foco, campo, valor_ant, valor_novo, usuario, alterado_em)
                        VALUES (?,?,?,?,?,?)
                    """, (id_foco, campo, ant, nov, usuario, agora))

        conn.commit()
    finally:
        conn.close()
    if request.headers.get("X-Requested-With") == "XMLHttpRequest":
        return jsonify({"ok": True})
    return redirect(url_for("foco_detalhe", id_foco=id_foco))

@app.route("/notificacoes/foco/<id_foco>/status", methods=["POST"])
@login_required
@nivel_min("operador")
def foco_status_rapido(id_foco):
    novo = request.json.get("status") if request.is_json else request.form.get("status")
    if novo not in STATUS_OPCOES + [None]:
        return jsonify({"erro": "Status invÃ¡lido"}), 400
    conn = get_db()
    try:
        ant_row = conn.execute("SELECT status_notificacao FROM focos_positivos WHERE id_foco=?", (id_foco,)).fetchone()
        ant = ant_row[0] if ant_row else None
        conn.execute("UPDATE focos_positivos SET status_notificacao=? WHERE id_foco=?", (novo, id_foco))
        if str(ant or "") != str(novo or ""):
            conn.execute("""INSERT INTO focos_historico (id_foco,campo,valor_ant,valor_novo,usuario,alterado_em)
                            VALUES (?,?,?,?,?,?)""",
                         (id_foco, "status_notificacao", ant, novo,
                          session.get("nome","desconhecido"), datetime.now().isoformat()))
        conn.commit()
    finally:
        conn.close()
    return jsonify({"ok": True, "status": novo})

@app.route("/notificacoes/imprimir", methods=["POST"])
@login_required
@nivel_min("operador")
def imprimir():
    ids = request.form.getlist("ids")
    if not ids: return redirect(url_for("notificacoes"))
    conn = get_db()
    focos = []
    for id_foco in ids:
        row = conn.execute("""
            SELECT f.*, l.nome AS localidade_nome FROM focos_positivos f
            LEFT JOIN localidades l ON l.id_localidade=f.id_localidade
            WHERE f.id_foco=? AND f.gera_notificacao=1
        """, (id_foco,)).fetchone()
        if row: focos.append(dict(row))
    if not focos:
        conn.close(); return "Nenhum foco vÃ¡lido.", 400
    try:
        caminho = gerar_docx(focos)
    except Exception as e:
        conn.close(); return f"Erro ao gerar DOCX: {e}", 500
    for f in focos:
        conn.execute("UPDATE focos_positivos SET status_notificacao='impressa' WHERE id_foco=? AND COALESCE(status_notificacao,'pendente')='pendente'", (f["id_foco"],))
    conn.commit(); conn.close()
    return send_file(caminho, as_attachment=True,
                     download_name=f"notificacoes_{datetime.now().strftime('%Y%m%d_%H%M')}.docx")

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  ROTAS â€” PROCESSAR (ETL via upload)
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

#  API JSON
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

# Documento de notificacao

def formatar_data_br(data_iso):
    meses = ["janeiro","fevereiro","marÃ§o","abril","maio","junho",
             "julho","agosto","setembro","outubro","novembro","dezembro"]
    try:
        d = datetime.strptime(str(data_iso)[:10], "%Y-%m-%d")
        return f"{d.day} de {meses[d.month-1]} de {d.year}"
    except (ValueError, TypeError, AttributeError):
        return "______"

def remover_bordas(tabela):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for row in tabela.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for side in ["top","bottom","left","right","insideH","insideV"]:
                b = OxmlElement(f"w:{side}"); b.set(qn("w:val"), "none"); borders.append(b)
            tcPr.append(borders)

def gerar_via(doc, foco, modelo):
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    def par(texto="", bold=False, italic=False, size=11,
            align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.space_after  = Pt(sa)
        if texto:
            r = p.add_run(texto); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
        return p

    def par_mixed(partes, align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=6):
        p = doc.add_paragraph(); p.alignment = align
        p.paragraph_format.space_before = Pt(sb); p.paragraph_format.space_after = Pt(sa)
        for texto, bold, italic, size in partes:
            r = p.add_run(texto); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
        return p

    cab = doc.add_table(rows=1, cols=3); cab.style = "Table Grid"; remover_bordas(cab)

    # Coluna esquerda â€” Logo da prefeitura
    cl = cab.rows[0].cells[0]; cl.width = Cm(3.5)
    pl = cl.paragraphs[0]; pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_pref = os.path.join(BASE_DIR, "static", "img", "logo_prefeitura.png")
    try: pl.add_run().add_picture(logo_pref, width=Cm(3))
    except Exception: pl.add_run("[LOGO PREFEITURA]")

    # Coluna central â€” TÃ­tulos
    ct = cab.rows[0].cells[1]
    for i, linha in enumerate(modelo.get("CABECALHO","").split("\n")):
        p = ct.paragraphs[0] if i == 0 else ct.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(linha); r.bold = (i==0); r.font.size = Pt(10 if i==0 else 9)

    # Coluna direita â€” Logo do Setor de Endemias
    cr = cab.rows[0].cells[2]; cr.width = Cm(3.5)
    pr = cr.paragraphs[0]; pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_end = os.path.join(BASE_DIR, "static", "img", "logo_endemias.png")
    try: pr.add_run().add_picture(logo_end, width=Cm(3))
    except Exception: pr.add_run("[LOGO ENDEMIAS]")

    # Centralizar a tabela na pÃ¡gina
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OxmlElement
    tblPr = cab._tbl.tblPr
    jc = _OxmlElement("w:jc"); jc.set(_qn("w:val"), "center")
    tblPr.append(jc)

    doc.add_paragraph()
    par(f"Almirante TamandarÃ© (PR), ______/______ de {date.today().year}.",
        align=WD_ALIGN_PARAGRAPH.RIGHT, size=10, sa=8)
    par(modelo.get("TITULO","COMUNICADO / NOTIFICAÃ‡ÃƒO"), bold=True, size=13,
        align=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
    par(modelo.get("SAUDACAO","Prezado(a) Senhor(a) PROPRIETÃRIO/RESPONSÃVEL"), bold=True, size=11, sa=8)

    end_fmt  = f"{foco.get('logradouro') or ''}, {foco.get('numero') or 's/n'}".strip(", ")
    loc_fmt  = foco.get("localidade_nome") or foco.get("localidade") or ""
    qrt_fmt  = f"QuarteirÃ£o {foco.get('quarteirao')}" if foco.get("quarteirao") else ""
    loc_linha = " â€” ".join(filter(None, [loc_fmt, qrt_fmt]))

    corpo = modelo.get("CORPO","").replace("{endereco}", end_fmt)\
                                  .replace("{localidade}", loc_linha)\
                                  .replace("{data_visita}", formatar_data_br(foco.get("data")))
    partes = []
    marcador = "Aedes aegypti"
    while marcador in corpo:
        idx = corpo.index(marcador)
        if corpo[:idx]: partes.append((corpo[:idx], False, False, 11))
        partes.append((marcador, False, True, 11)); corpo = corpo[idx+len(marcador):]
    if corpo: partes.append((corpo, False, False, 11))
    par_mixed(partes, sa=8)

    par(modelo.get("AVISO",""), bold=True, size=11, sa=8)
    par(modelo.get("CONTATO",""), size=10, sa=10)
    campos = [
        ("LOCALIDADE",    loc_linha or "___________________________"),
        ("ENDEREÃ‡O",      end_fmt   or "___________________________"),
        ("MORADOR",       foco.get("nome_morador") or "___________________________"),
        ("DEPÃ“SITO(S)",   foco.get("depositos")    or "___________________________"),
        ("AGENTE(S)",     foco.get("agentes")      or "___________________________"),
    ]
    if foco.get("observacoes"):
        campos.append(("OBSERVAÃ‡Ã•ES", foco["observacoes"]))
    for label, valor in campos:
        par_mixed([(f"â€¢ {label}: ", True, True, 11), (valor, False, True, 11)], sa=3)

    doc.add_paragraph()
    ass = doc.add_table(rows=2, cols=2); ass.style = "Table Grid"; remover_bordas(ass)
    for i, txt in enumerate(["_"*35, "_"*35]):
        c = ass.rows[0].cells[i]; c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraphs[0].add_run(txt)
    for i, txt in enumerate(["VigilÃ¢ncia Ambiental / Setor de Endemias","ProprietÃ¡rio / ResponsÃ¡vel"]):
        c = ass.rows[1].cells[i]; c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraphs[0].add_run(txt).bold = True

    doc.add_paragraph()
    par(modelo.get("RODAPE",""), size=8, align=WD_ALIGN_PARAGRAPH.CENTER, sa=0)
    if foco.get("codigo"):
        par(f"NÂº da notificaÃ§Ã£o: {foco['codigo']}", size=7,
            align=WD_ALIGN_PARAGRAPH.CENTER, sa=0)

def gerar_docx(focos):
    from docx import Document; from docx.shared import Cm
    modelo = ler_modelo(); doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(1.5)
        section.left_margin = section.right_margin  = Cm(2)
    for i, foco in enumerate(focos):
        gerar_via(doc, foco, modelo)
        if i < len(focos)-1: doc.add_page_break()
    caminho = os.path.join(SAIDA_DIR, f"notificacoes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    doc.save(caminho); return caminho

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  MAPA
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•


def _focos_para_impressao(ids):
    """Retorna lista de dicts de focos para impressÃ£o, na ordem dos IDs."""
    conn = get_db()
    focos = []
    for id_foco in ids:
        row = conn.execute("""
            SELECT f.*, l.nome AS localidade_nome
            FROM focos_positivos f
            LEFT JOIN localidades l ON l.id_localidade = f.id_localidade
            WHERE f.id_foco = ? AND f.gera_notificacao = 1
        """, (id_foco,)).fetchone()
        if row:
            focos.append(dict(row))
    conn.close()
    return focos

@app.route("/notificacoes/foco/<id_foco>/imprimir-html")
@login_required
@nivel_min("operador")
def imprimir_html_single(id_foco):
    focos = _focos_para_impressao([id_foco])
    if not focos:
        abort(404)
    # Marcar como impressa se ainda pendente
    conn = get_db()
    conn.execute("""UPDATE focos_positivos SET status_notificacao='impressa'
                    WHERE id_foco=? AND COALESCE(status_notificacao,'pendente')='pendente'""", (id_foco,))
    conn.commit(); conn.close()
    return render_template("notificacao_print.html", focos=focos, auto_print=True, modelo=type("M", (), ler_modelo())())

@app.route("/notificacoes/imprimir-html", methods=["POST"])
@login_required
@nivel_min("operador")
def imprimir_html_lote():
    ids = request.form.getlist("ids")
    if not ids:
        return redirect(url_for("notificacoes"))
    focos = _focos_para_impressao(ids)
    if not focos:
        return "Nenhum foco vÃ¡lido.", 400
    conn = get_db()
    for f in focos:
        conn.execute("""UPDATE focos_positivos SET status_notificacao='impressa'
                        WHERE id_foco=? AND COALESCE(status_notificacao,'pendente')='pendente'""",
                     (f["id_foco"],))
    conn.commit(); conn.close()
    return render_template("notificacao_print.html", focos=focos, auto_print=True, modelo=type("M", (), ler_modelo())())

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  ERROR HANDLERS
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

@app.errorhandler(404)
def err404(e): return render_template("404.html"), 404

@app.errorhandler(500)
def err500(e): return render_template("500.html"), 500

@app.errorhandler(403)
def err403(e): return render_template("403.html"), 403

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  MAIN
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

if __name__ == "__main__":
    import socket
    hostname = socket.gethostname()
    try:    ip = socket.gethostbyname(hostname)
    except OSError: ip = "127.0.0.1"

    print("=" * 54)
    print("  ENDEMIAS â€” Sistema de GestÃ£o Integrado  v3")
    print("  Setor de Endemias Â· Almirante TamandarÃ©-PR")
    print("=" * 54)
    print(f"\n  Banco de dados: {DB_PATH}")
    print(f"\n  Acesse no navegador:")
    print(f"    Este computador : http://localhost:5000")
    print(f"    Rede local      : http://{ip}:5000")
    print(f"\n  Para encerrar: Ctrl+C ou feche esta janela")
    print("=" * 54 + "\n")

    app.run(debug=False, host="0.0.0.0", port=5000, threaded=True)

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
